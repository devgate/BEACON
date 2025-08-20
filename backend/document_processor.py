"""
Document Processing Service
Handles text extraction from various file formats and chunking strategies
"""

import os
import logging
from typing import List, Dict, Any, Optional, Tuple
import mimetypes
from pathlib import Path

# Document processing imports
import PyPDF2
from docx import Document  # Enable DOCX support
import openpyxl
import json
import csv
from io import StringIO

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """
    Comprehensive document processor for multiple file formats
    """
    
    def __init__(self):
        """Initialize document processor"""
        self.supported_formats = {
            '.pdf': self._extract_pdf_text,
            '.txt': self._extract_txt_text,
            '.docx': self._extract_docx_text,  # Enable DOCX support
            '.doc': self._extract_docx_text,   # Enable DOC support
            '.xlsx': self._extract_xlsx_text,  # Enable Excel support
            '.xls': self._extract_xlsx_text,   # Enable Excel support
            '.csv': self._extract_csv_text,
            '.json': self._extract_json_text,
            '.md': self._extract_txt_text,  # Markdown as text
            '.rtf': self._extract_txt_text,  # RTF as text
        }
    
    def extract_text(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text from a document file
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")
            
            # Get file extension
            file_extension = Path(file_path).suffix.lower()
            
            if file_extension not in self.supported_formats:
                raise ValueError(f"Unsupported file format: {file_extension}")
            
            # Extract text using appropriate method
            extractor = self.supported_formats[file_extension]
            text, metadata = extractor(file_path)
            
            # Add common metadata
            file_stats = os.stat(file_path)
            metadata.update({
                'file_path': file_path,
                'file_size': file_stats.st_size,
                'file_extension': file_extension,
                'extraction_method': extractor.__name__
            })
            
            logger.info(f"Extracted {len(text)} characters from {Path(file_path).name}")
            return text, metadata
            
        except Exception as e:
            logger.error(f"Failed to extract text from {file_path}: {str(e)}")
            raise
    
    def _extract_pdf_text(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from PDF file"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text_parts = []
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_parts.append(page_text)
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num + 1}: {str(e)}")
                
                extracted_text = '\n\n'.join(text_parts)
                
                metadata = {
                    'total_pages': len(pdf_reader.pages),
                    'pages_processed': len(text_parts),
                    'has_metadata': bool(pdf_reader.metadata),
                    'pdf_metadata': dict(pdf_reader.metadata) if pdf_reader.metadata else {}
                }
                
                return extracted_text, metadata
                
        except Exception as e:
            logger.error(f"PDF extraction failed: {str(e)}")
            raise
    
    def _extract_txt_text(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from text file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                    
                    metadata = {
                        'encoding_used': encoding,
                        'line_count': text.count('\n') + 1
                    }
                    
                    return text, metadata
                    
                except UnicodeDecodeError:
                    continue
            
            raise ValueError("Unable to decode text file with any supported encoding")
            
        except Exception as e:
            logger.error(f"Text extraction failed: {str(e)}")
            raise
    
    def _extract_docx_text(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_texts.append(' | '.join(row_text))
            
            if table_texts:
                text_parts.extend(table_texts)
            
            extracted_text = '\n\n'.join(text_parts)
            
            metadata = {
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables),
                'has_tables': len(doc.tables) > 0,
                'estimated_words': len(extracted_text.split()) if extracted_text else 0
            }
            
            return extracted_text, metadata
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {str(e)}")
            raise
    
    def _extract_xlsx_text(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from Excel file"""
        try:
            workbook = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
            text_parts = []
            worksheet_info = []
            
            for sheet_name in workbook.sheetnames:
                worksheet = workbook[sheet_name]
                sheet_text_parts = []
                
                # Add sheet header
                sheet_text_parts.append(f"=== {sheet_name} ===")
                
                row_count = 0
                for row in worksheet.iter_rows():
                    row_values = []
                    for cell in row:
                        if cell.value is not None:
                            row_values.append(str(cell.value))
                    
                    if row_values and any(val.strip() for val in row_values):
                        sheet_text_parts.append(' | '.join(row_values))
                        row_count += 1
                    
                    # Limit rows for very large files
                    if row_count > 1000:
                        sheet_text_parts.append("... (truncated, too many rows)")
                        break
                
                if sheet_text_parts and len(sheet_text_parts) > 1:  # More than just header
                    text_parts.extend(sheet_text_parts)
                    worksheet_info.append({
                        'name': sheet_name,
                        'rows_processed': row_count
                    })
            
            extracted_text = '\n\n'.join(text_parts)
            
            metadata = {
                'worksheet_count': len(workbook.sheetnames),
                'worksheets_processed': len(worksheet_info),
                'worksheet_info': worksheet_info
            }
            
            workbook.close()
            return extracted_text, metadata
            
        except Exception as e:
            logger.error(f"Excel extraction failed: {str(e)}")
            raise
    
    def _extract_csv_text(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from CSV file"""
        try:
            # Try different encodings and delimiters
            encodings = ['utf-8', 'latin-1', 'cp1252']
            delimiters = [',', ';', '\t', '|']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        # Detect delimiter
                        sample = file.read(1024)
                        file.seek(0)
                        
                        sniffer = csv.Sniffer()
                        delimiter = sniffer.sniff(sample).delimiter
                        
                        # Read CSV
                        csv_reader = csv.reader(file, delimiter=delimiter)
                        rows = []
                        
                        for row_num, row in enumerate(csv_reader):
                            if row and any(cell.strip() for cell in row):  # Skip empty rows
                                rows.append(' | '.join(row))
                            
                            # Limit rows for very large files
                            if row_num > 10000:
                                logger.warning(f"CSV file too large, limiting to first 10000 rows")
                                break
                        
                        extracted_text = '\n'.join(rows)
                        
                        metadata = {
                            'encoding_used': encoding,
                            'delimiter_used': delimiter,
                            'row_count': len(rows),
                            'estimated_columns': len(rows[0].split(' | ')) if rows else 0
                        }
                        
                        return extracted_text, metadata
                        
                except UnicodeDecodeError:
                    continue
            
            raise ValueError("Unable to decode CSV file with any supported encoding")
            
        except Exception as e:
            logger.error(f"CSV extraction failed: {str(e)}")
            raise
    
    def _extract_json_text(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from JSON file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
            
            # Convert JSON to readable text
            def json_to_text(obj, level=0):
                """Convert JSON object to readable text"""
                lines = []
                indent = "  " * level
                
                if isinstance(obj, dict):
                    for key, value in obj.items():
                        if isinstance(value, (dict, list)):
                            lines.append(f"{indent}{key}:")
                            lines.extend(json_to_text(value, level + 1))
                        else:
                            lines.append(f"{indent}{key}: {str(value)}")
                elif isinstance(obj, list):
                    for i, item in enumerate(obj):
                        if isinstance(item, (dict, list)):
                            lines.append(f"{indent}[{i}]:")
                            lines.extend(json_to_text(item, level + 1))
                        else:
                            lines.append(f"{indent}[{i}]: {str(item)}")
                else:
                    lines.append(f"{indent}{str(obj)}")
                
                return lines
            
            text_lines = json_to_text(data)
            extracted_text = '\n'.join(text_lines)
            
            metadata = {
                'json_type': type(data).__name__,
                'estimated_size': len(str(data)),
                'line_count': len(text_lines)
            }
            
            return extracted_text, metadata
            
        except Exception as e:
            logger.error(f"JSON extraction failed: {str(e)}")
            raise
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats"""
        return list(self.supported_formats.keys())
    
    def is_supported_format(self, file_path: str) -> bool:
        """Check if file format is supported"""
        file_extension = Path(file_path).suffix.lower()
        return file_extension in self.supported_formats
    
    def get_file_info(self, file_path: str) -> Dict[str, Any]:
        """Get basic file information"""
        try:
            if not os.path.exists(file_path):
                return {"exists": False}
            
            file_stats = os.stat(file_path)
            path_obj = Path(file_path)
            
            # Detect MIME type
            mime_type, _ = mimetypes.guess_type(file_path)
            
            return {
                "exists": True,
                "name": path_obj.name,
                "extension": path_obj.suffix.lower(),
                "size": file_stats.st_size,
                "size_mb": round(file_stats.st_size / (1024 * 1024), 2),
                "mime_type": mime_type,
                "is_supported": self.is_supported_format(file_path),
                "modified_time": file_stats.st_mtime
            }
            
        except Exception as e:
            logger.error(f"Failed to get file info for {file_path}: {str(e)}")
            return {"exists": False, "error": str(e)}